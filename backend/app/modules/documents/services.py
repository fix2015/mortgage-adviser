import io
import json as _json
import re
import uuid
from typing import Optional

import boto3
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.config import settings
from app.modules.documents.models import Document, DocumentStatus


ALLOWED_FILE_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "text/plain": "txt",
    "text/csv": "csv",
    "image/jpeg": "jpg",
    "image/png": "png",
    "image/jpg": "jpg",
}

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "csv", "jpg", "jpeg", "png"}


def get_s3_client():
    return boto3.client(
        "s3",
        aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
        region_name=settings.AWS_S3_REGION,
    )


def upload_to_s3(file_content: bytes, s3_key: str, content_type: str) -> str:
    s3 = get_s3_client()
    s3.put_object(
        Bucket=settings.AWS_S3_BUCKET,
        Key=s3_key,
        Body=file_content,
        ContentType=content_type,
    )
    return f"{settings.S3_BASE_URL}/{s3_key}"


def extract_text_from_pdf(file_content: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_content))
    text_parts = []
    for page in reader.pages:
        try:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
        except Exception:
            continue
    result = "\n\n".join(text_parts)
    if not result.strip():
        # Fallback: try extracting with layout mode
        try:
            reader2 = PdfReader(io.BytesIO(file_content))
            for page in reader2.pages:
                page_text = page.extract_text(extraction_mode="layout")
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            result = "\n\n".join(text_parts)
        except Exception:
            pass
    return (
        result
        if result.strip()
        else "PDF content could not be extracted automatically."
    )


def extract_text_from_docx(file_content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_content))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_content: bytes) -> str:
    return file_content.decode("utf-8", errors="replace")


def extract_text_from_csv(file_content: bytes) -> str:
    return file_content.decode("utf-8", errors="replace")


def extract_text_from_image(file_content: bytes) -> str:
    return "Image document (passport, driving licence, or ID scan). Content requires visual inspection."


def extract_text(file_content: bytes, file_type: str) -> str:
    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "doc": extract_text_from_docx,
        "txt": extract_text_from_txt,
        "csv": extract_text_from_csv,
        "jpg": extract_text_from_image,
        "jpeg": extract_text_from_image,
        "png": extract_text_from_image,
    }
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    return extractor(file_content)


def classify_and_extract(db: Session, document: Document) -> Document:
    """Use OpenAI to classify the document type and extract key structured data for mortgage applications."""
    if not document.extracted_text:
        return document

    # Truncate text to avoid token limits
    text_sample = document.extracted_text[:4000]

    prompt = (
        "Analyze this document for a UK mortgage application and return ONLY valid JSON:\n"
        '{"document_type": "passport|payslip|bank_statement|tax_return|contract|utility_bill|credit_report|company_accounts|other", '
        '"checklist_category": "id|address|immigration|income|bank_statements|deposit|credit_report|employment|company_accounts|payslips|tax_returns|other", '
        '"key_data": {"total_amount": "£X", "date": "YYYY-MM-DD", "from": "entity", '
        '"to": "entity", "description": "brief", "employer": "name if payslip/contract", '
        '"income_amount": "£X if income doc"}, '
        '"flags": ["any anomalies or notable items relevant to mortgage application"]}\n\n'
        f"Document text:\n{text_sample}"
    )

    try:
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0.2,
        )

        raw = response.choices[0].message.content.strip()
        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)

        data = _json.loads(raw)

        # Save classification
        doc_type = data.get("document_type", "other")
        valid_types = {
            "passport",
            "payslip",
            "bank_statement",
            "tax_return",
            "contract",
            "utility_bill",
            "credit_report",
            "company_accounts",
            "other",
        }
        document.document_type = doc_type if doc_type in valid_types else "other"

        checklist_cat = data.get("checklist_category", "other")
        valid_categories = {
            "id",
            "address",
            "immigration",
            "income",
            "bank_statements",
            "deposit",
            "credit_report",
            "employment",
            "company_accounts",
            "payslips",
            "tax_returns",
            "other",
        }
        document.checklist_category = (
            checklist_cat if checklist_cat in valid_categories else "other"
        )

        document.structured_data = _json.dumps(data)

    except Exception:
        # If classification fails, set defaults -- don't break the upload flow
        document.document_type = "other"
        document.checklist_category = "other"
        document.structured_data = None

    db.commit()
    db.refresh(document)
    return document


def create_document(
    db: Session,
    consultation_id: int,
    user_id: int,
    filename: str,
    file_content: bytes,
    content_type: str,
    file_size: int,
) -> Document:
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"File type .{extension} is not supported. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Check for duplicate: same filename + same file size in this consultation
    existing = (
        db.query(Document)
        .filter(
            Document.consultation_id == consultation_id,
            Document.filename == filename,
            Document.file_size == file_size,
        )
        .first()
    )
    if existing:
        raise ValueError(f"File '{filename}' has already been uploaded")

    s3_key = f"{settings.AWS_S3_PREFIX}/documents/{user_id}/{consultation_id}/{uuid.uuid4().hex}_{filename}"

    upload_to_s3(file_content, s3_key, content_type)

    document = Document(
        consultation_id=consultation_id,
        user_id=user_id,
        filename=filename,
        s3_key=s3_key,
        file_type=extension,
        file_size=file_size,
        status=DocumentStatus.UPLOADED,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


def process_document(db: Session, document_id: int) -> Document:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise ValueError(f"Document {document_id} not found")

    document.status = DocumentStatus.PROCESSING
    db.commit()

    try:
        s3 = get_s3_client()
        response = s3.get_object(Bucket=settings.AWS_S3_BUCKET, Key=document.s3_key)
        file_content = response["Body"].read()

        extracted_text = extract_text(file_content, document.file_type)
        document.extracted_text = extracted_text
        document.status = DocumentStatus.PROCESSED
    except Exception as e:
        document.status = DocumentStatus.ERROR
        document.error_message = str(e)

    db.commit()
    db.refresh(document)
    return document


def process_document_inline(
    db: Session, document: Document, file_content: bytes
) -> Document:
    """Process document text extraction inline (without Celery)."""
    document.status = DocumentStatus.PROCESSING
    db.commit()

    try:
        extracted_text = extract_text(file_content, document.file_type)
        document.extracted_text = extracted_text
        document.status = DocumentStatus.PROCESSED
        db.commit()
        db.refresh(document)
    except Exception as e:
        document.status = DocumentStatus.ERROR
        document.error_message = str(e)
        db.commit()
        db.refresh(document)
        return document

    # Classify document type -- separate try/catch so extraction failures
    # don't prevent the document from being marked as processed
    try:
        classify_and_extract(db, document)
    except Exception:
        pass  # Classification is optional, don't fail the whole upload

    return document


def get_consultation_documents(
    db: Session, consultation_id: int, skip: int = 0, limit: int = 50
) -> tuple[list[Document], int]:
    query = db.query(Document).filter(Document.consultation_id == consultation_id)
    total = query.count()
    documents = (
        query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()
    )
    return documents, total


def get_document_by_id(
    db: Session, document_id: int, user_id: Optional[int] = None
) -> Optional[Document]:
    query = db.query(Document).filter(Document.id == document_id)
    if user_id:
        query = query.filter(Document.user_id == user_id)
    return query.first()


def delete_document(db: Session, document: Document) -> None:
    try:
        s3 = get_s3_client()
        s3.delete_object(Bucket=settings.AWS_S3_BUCKET, Key=document.s3_key)
    except Exception:
        pass  # Continue even if S3 delete fails
    db.delete(document)
    db.commit()
